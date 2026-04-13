"""DOCX ingest via python-docx.

Walks the document body in ORDER so paragraph → table → paragraph stays in
that sequence.  Earlier versions iterated `doc.paragraphs` then `doc.tables`
separately, which silently re-ordered mixed layouts.

Output is markdown-ish: headings get `#` prefixes based on the Word style,
body paragraphs are blank-line-separated, tables render as markdown tables.

Trade-off: python-docx doesn't preserve nested lists, text-box content, or
drawings.  Pandoc would do better but requires a system binary; for v1 we
optimise for zero-dependency import.
"""
from __future__ import annotations

import io

from docx import Document
from docx.document import Document as _Doc
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


def _iter_body_blocks(doc: _Doc):
    """Yield paragraphs and tables in document order.

    python-docx doesn't expose a body iterator; we walk the underlying XML
    children of `<w:body>` and dispatch on tag name.
    """
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)
        # other tags (sectPr, etc.) are metadata — skip silently


def _render_paragraph(p: Paragraph) -> str | None:
    text = p.text.strip()
    if not text:
        return None
    style = (p.style.name or "").lower()
    if style.startswith("heading 1"):
        return f"# {text}"
    if style.startswith("heading 2"):
        return f"## {text}"
    if style.startswith("heading 3"):
        return f"### {text}"
    if style.startswith("heading"):
        return f"#### {text}"
    return text


def _render_table(table: Table) -> str | None:
    if not table.rows:
        return None
    rows = [
        ["".join(cell.text.splitlines()).strip() for cell in row.cells]
        for row in table.rows
    ]
    if not rows or not any(any(c for c in r) for r in rows):
        return None

    width = max(len(r) for r in rows)
    lines: list[str] = []
    header = rows[0] + [""] * (width - len(rows[0]))
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join(["---"] * width) + " |")
    for row in rows[1:]:
        padded = row + [""] * (width - len(row))
        lines.append("| " + " | ".join(padded) + " |")
    return "\n".join(lines)


def parse(content: bytes) -> dict:
    doc = Document(io.BytesIO(content))

    parts: list[str] = []
    for block in _iter_body_blocks(doc):
        if isinstance(block, Paragraph):
            rendered = _render_paragraph(block)
        else:  # Table
            rendered = _render_table(block)
        if rendered is not None:
            parts.append(rendered)

    text = "\n\n".join(parts)
    return {"text": text, "source_type": "docx", "warnings": []}
