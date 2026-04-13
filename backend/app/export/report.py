"""Writing-process report export in markdown + DOCX.

Renders the provenance report (from provenance.service.build_report) into a
human-readable format suitable for attaching to an academic submission or
sending to an editor.
"""
from __future__ import annotations

import io
from datetime import datetime, timezone

from docx import Document


def _format_time(ms: int) -> str:
    return (
        datetime.fromtimestamp(ms / 1000, tz=timezone.utc)
        .astimezone()
        .strftime("%Y-%m-%d %H:%M:%S")
    )


def _duration(started: int, ended: int | None) -> str:
    if ended is None:
        return "(active)"
    total_s = (ended - started) // 1000
    if total_s < 60:
        return f"{total_s}s"
    m, s = divmod(total_s, 60)
    if m < 60:
        return f"{m}m {s}s"
    h, m = divmod(m, 60)
    return f"{h}h {m}m"


def report_to_markdown(report: dict) -> str:
    lines: list[str] = []
    lines.append(f"# Writing Process Report — {report.get('document_title', 'Document')}")
    lines.append("")

    integrity = report.get("integrity", {})
    mark = "✓" if integrity.get("valid") else "⚠"
    lines.append(f"**Chain integrity:** {mark} {'verified' if integrity.get('valid') else 'broken'}")
    lines.append(
        f"**Sessions verified:** {integrity.get('sessions_verified', 0)} · "
        f"**Total events:** {report.get('total_events', 0)}"
    )
    lines.append("")

    # Authorship
    a = report.get("authorship", {})
    lines.append("## Authorship")
    lines.append("")
    lines.append(f"- Typed: {a.get('typed_chars', 0)} chars ({a.get('typed_pct', 0)}%)")
    lines.append(f"- Pasted: {a.get('pasted_chars', 0)} chars ({a.get('pasted_pct', 0)}%)")
    lines.append(f"- AI-assisted: {a.get('ai_assisted_chars', 0)} chars ({a.get('ai_assisted_pct', 0)}%)")
    lines.append("")

    # Sessions
    lines.append("## Sessions")
    lines.append("")
    for s in report.get("sessions", []):
        status = "✓" if s.get("valid") else "✗"
        lines.append(
            f"- {status} {_format_time(s['started_at'])} · "
            f"{_duration(s['started_at'], s.get('ended_at'))} · "
            f"{s.get('events', 0)} events · "
            f"hash `{(s.get('final_hash') or '—')[:16]}`"
        )
    lines.append("")

    # Timeline
    lines.append("## Timeline")
    lines.append("")
    for e in report.get("timeline", []):
        ts = _format_time(e["timestamp"])
        lines.append(f"- {ts} · **{e['event_type']}** · {e['summary']}")

    return "\n".join(lines)


def report_to_docx(report: dict) -> bytes:
    md = report_to_markdown(report)
    doc = Document()
    for block in md.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("# "):
            doc.add_heading(block[2:], level=1)
        elif block.startswith("## "):
            doc.add_heading(block[3:], level=2)
        elif block.startswith("### "):
            doc.add_heading(block[4:], level=3)
        else:
            for line in block.split("\n"):
                if line.startswith("- "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                else:
                    doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
