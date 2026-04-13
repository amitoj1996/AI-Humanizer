"""DOCX export via python-docx.

Straightforward deterministic template: one paragraph per blank-line-
separated block, with markdown-ish `#`/`##`/`###` prefixes upgraded to Word
heading styles.  This is the "we control the template" export — Pandoc-based
style preservation is a future add-on.
"""
from __future__ import annotations

import io

from docx import Document


def export(text: str, title: str) -> bytes:
    doc = Document()

    core = doc.core_properties
    core.title = title

    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        level = 0
        while block.startswith("#"):
            level += 1
            block = block[1:]
        block = block.lstrip()
        if level and level <= 4 and block:
            doc.add_heading(block, level=min(level, 4))
        elif block:
            doc.add_paragraph(block)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
