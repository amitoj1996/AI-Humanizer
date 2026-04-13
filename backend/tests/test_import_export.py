"""Import + export round-trip tests.

Builds tiny real DOCX/PDF files in-memory so we exercise the actual parsers
(not just stubs).  PDF uses pymupdf to synthesise a 1-page PDF; DOCX uses
python-docx.
"""
from __future__ import annotations

import io

import pymupdf
from docx import Document


def _make_docx(paragraphs: list[tuple[str, str]]) -> bytes:
    """Build a DOCX with (style, text) pairs."""
    doc = Document()
    for style, text in paragraphs:
        if style.lower().startswith("heading"):
            level = int(style.split()[-1]) if style.split()[-1].isdigit() else 1
            doc.add_heading(text, level=level)
        else:
            doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_pdf(text: str) -> bytes:
    """Build a single-page PDF using pymupdf (already a transitive dep)."""
    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((72, 72), text, fontsize=12)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _create_project(client) -> str:
    return client.post("/api/projects", json={"name": "Import Test"}).json()["id"]


# ---------------------------------------------------------------------------
# Import
# ---------------------------------------------------------------------------
def test_import_txt(client):
    project_id = _create_project(client)
    files = {"file": ("notes.txt", b"hello world\nline two", "text/plain")}
    r = client.post(
        "/api/documents/import",
        data={"project_id": project_id},
        files=files,
    )
    assert r.status_code == 200
    result = r.json()
    assert result["source_type"] == "txt"
    assert result["title"] == "notes"
    assert result["char_count"] > 0

    # Document should exist with the imported content
    doc = client.get(f"/api/documents/{result['document_id']}").json()
    assert doc["current_revision_id"]
    rev = client.get(
        f"/api/documents/{result['document_id']}/revisions/{doc['current_revision_id']}"
    ).json()
    assert "hello world" in rev["content"]


def test_import_md(client):
    project_id = _create_project(client)
    md = b"# Heading\n\nparagraph"
    r = client.post(
        "/api/documents/import",
        data={"project_id": project_id},
        files={"file": ("notes.md", md, "text/markdown")},
    )
    assert r.status_code == 200
    assert r.json()["source_type"] == "md"


def test_import_docx_preserves_headings(client):
    project_id = _create_project(client)
    content = _make_docx(
        [
            ("Heading 1", "Top heading"),
            ("Normal", "Body text here."),
            ("Heading 2", "Sub heading"),
            ("Normal", "More body."),
        ]
    )
    r = client.post(
        "/api/documents/import",
        data={"project_id": project_id},
        files={
            "file": (
                "essay.docx",
                content,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert r.status_code == 200
    result = r.json()
    assert result["source_type"] == "docx"

    doc = client.get(f"/api/documents/{result['document_id']}").json()
    rev = client.get(
        f"/api/documents/{result['document_id']}/revisions/{doc['current_revision_id']}"
    ).json()
    assert "# Top heading" in rev["content"]
    assert "## Sub heading" in rev["content"]
    assert "Body text here." in rev["content"]


def test_import_pdf(client):
    project_id = _create_project(client)
    pdf_bytes = _make_pdf("Hello from a PDF. Line two.")
    r = client.post(
        "/api/documents/import",
        data={"project_id": project_id},
        files={"file": ("paper.pdf", pdf_bytes, "application/pdf")},
    )
    assert r.status_code == 200
    result = r.json()
    assert result["source_type"] == "pdf"
    assert result["char_count"] > 0


def test_import_rejects_unknown_extension(client):
    project_id = _create_project(client)
    r = client.post(
        "/api/documents/import",
        data={"project_id": project_id},
        files={"file": ("weird.xyz", b"data", "application/octet-stream")},
    )
    assert r.status_code == 400


def test_import_rejects_missing_project(client):
    r = client.post(
        "/api/documents/import",
        data={"project_id": "does-not-exist"},
        files={"file": ("x.txt", b"hi", "text/plain")},
    )
    assert r.status_code == 404


def test_import_rejects_empty_file(client):
    project_id = _create_project(client)
    r = client.post(
        "/api/documents/import",
        data={"project_id": project_id},
        files={"file": ("empty.txt", b"", "text/plain")},
    )
    assert r.status_code == 400


# ---------------------------------------------------------------------------
# Export
# ---------------------------------------------------------------------------
def _make_doc_with_content(client, content: str, title: str = "Essay") -> str:
    p_id = _create_project(client)
    doc = client.post(
        "/api/documents",
        json={"project_id": p_id, "title": title, "initial_content": content},
    ).json()
    return doc["id"]


def test_export_md(client):
    doc_id = _make_doc_with_content(client, "# Title\n\nBody text.")
    r = client.get(f"/api/documents/{doc_id}/export?format=md")
    assert r.status_code == 200
    assert r.headers["content-type"].startswith("text/markdown")
    assert "# Title" in r.text
    assert "attachment" in r.headers["content-disposition"]
    assert ".md" in r.headers["content-disposition"]


def test_export_txt(client):
    doc_id = _make_doc_with_content(client, "plain text only")
    r = client.get(f"/api/documents/{doc_id}/export?format=txt")
    assert r.status_code == 200
    assert r.text == "plain text only"


def test_export_docx_is_valid(client):
    """Round-trip: export DOCX then parse it back with python-docx."""
    doc_id = _make_doc_with_content(
        client, "# Chapter One\n\nOpening paragraph.\n\n## Section\n\nMore text."
    )
    r = client.get(f"/api/documents/{doc_id}/export?format=docx")
    assert r.status_code == 200
    assert r.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    parsed = Document(io.BytesIO(r.content))
    paragraphs = [p.text for p in parsed.paragraphs if p.text.strip()]
    assert "Chapter One" in paragraphs
    assert "Opening paragraph." in paragraphs


def test_export_rejects_unknown_format(client):
    doc_id = _make_doc_with_content(client, "hi")
    r = client.get(f"/api/documents/{doc_id}/export?format=bogus")
    assert r.status_code == 400


def test_export_rejects_empty_doc(client):
    p_id = _create_project(client)
    doc = client.post(
        "/api/documents", json={"project_id": p_id, "title": "Empty"}
    ).json()
    r = client.get(f"/api/documents/{doc['id']}/export?format=md")
    assert r.status_code == 400


# ---------------------------------------------------------------------------
# Provenance report export
# ---------------------------------------------------------------------------
def test_provenance_report_export_md(client):
    doc_id = _make_doc_with_content(client, "hi")
    r = client.get(f"/api/documents/{doc_id}/provenance/export?format=md")
    assert r.status_code == 200
    assert r.headers["content-type"].startswith("text/markdown")
    assert "Writing Process Report" in r.text


def test_provenance_report_export_docx(client):
    doc_id = _make_doc_with_content(client, "hi")
    r = client.get(f"/api/documents/{doc_id}/provenance/export?format=docx")
    assert r.status_code == 200
    assert r.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    # Validate it's really a DOCX
    parsed = Document(io.BytesIO(r.content))
    texts = [p.text for p in parsed.paragraphs]
    assert any("Writing Process Report" in t for t in texts)
